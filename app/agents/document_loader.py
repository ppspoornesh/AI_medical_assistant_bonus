from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_ollama import OllamaEmbeddings
from langchain_community.vectorstores import FAISS
from pypdf import PdfReader
from docx import Document as DocxDocument
import pandas as pd
from PIL import Image
import pytesseract
import os
from loguru import logger

class DocumentLoaderAgent:
    def __init__(self):
        logger.info("Initializing DocumentLoaderAgent – preparing for multi-format ingestion")
        # Using nomic-embed-text: optimized for semantic search
        self.embeddings = OllamaEmbeddings(model="nomic-embed-text")
        # Smaller chunks = better retrieval accuracy
        self.splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=100)
        self.vectorstore_path = "vectorstore/index"  # Persistent storage – survives restarts

    def load_documents(self, file_paths):
        logger.info(f"Loading {len(file_paths)} document(s): {file_paths}")
        existing_files = [f for f in file_paths if os.path.exists(f)]
        if not existing_files:
            logger.error("No valid files found – aborting load")
            return None

        docs = []
        for path in existing_files:
            try:
                if path.endswith('.pdf'):
                    reader = PdfReader(path)
                    text = "".join(page.extract_text() or "" for page in reader.pages)
                    docs.append(Document(page_content=text, metadata={"source": path, "type": "pdf"}))
                    logger.debug(f"Extracted text from PDF: {path}")

                elif path.endswith('.docx'):
                    doc = DocxDocument(path)
                    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                    docs.append(Document(page_content=text, metadata={"source": path, "type": "docx"}))
                    logger.debug(f"Parsed DOCX: {path}")

                elif path.endswith('.xlsx'):
                    df_dict = pd.read_excel(path, sheet_name=None)
                    text = "\n".join(sheet.to_string() for sheet in df_dict.values())
                    docs.append(Document(page_content=text, metadata={"source": path, "type": "xlsx"}))
                    logger.debug(f"Converted Excel sheets to text: {path}")

                elif path.endswith(('.png', '.jpg', '.jpeg')):
                    img = Image.open(path)
                    text = pytesseract.image_to_string(img)
                    docs.append(Document(page_content=text, metadata={"source": path, "type": "image"}))
                    logger.debug(f"OCR extracted from image: {path}")

            except Exception as e:
                logger.warning(f"Failed to process {path}: {e} – skipping file")
                continue

        if not docs:
            logger.error("All files failed to process")
            return None

        # Split into chunks for better RAG performance
        splits = self.splitter.split_documents(docs)
        logger.info(f"Split into {len(splits)} chunks")

        try:
            if os.path.exists(self.vectorstore_path):
                vectorstore = FAISS.load_local(
                    self.vectorstore_path, 
                    self.embeddings, 
                    allow_dangerous_deserialization=True
                )
                vectorstore.add_documents(splits)
                logger.info("Updated existing vectorstore")
            else:
                vectorstore = FAISS.from_documents(splits, self.embeddings)
                logger.info("Created new vectorstore")

            vectorstore.save_local(self.vectorstore_path)
            logger.success("Vectorstore saved to disk – persistent RAG ready!")
        except Exception as e:
            logger.error(f"Vectorstore creation failed: {e}")
            return None

        return vectorstore


if __name__ == "__main__":
    loader = DocumentLoaderAgent()
    file_paths = input("Enter file paths (comma-separated): ").strip()
    file_paths = [fp.strip() for fp in file_paths.split(",") if fp.strip()] or ["docs/sample_dataset/cmh-2022-0365.pdf"]
    vectorstore = loader.load_documents(file_paths)
    if vectorstore:
        print("Documents loaded and indexed successfully!")